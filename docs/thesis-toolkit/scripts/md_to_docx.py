"""
md_to_docx.py — Markdown dosyasini DOCX formatina donustur
Kullanim: python scripts/md_to_docx.py <kaynak.md> [hedef.docx]

Oncelik sirasi:
  1. python-docx varsa direkt docx uret
  2. pandoc varsa subprocess ile calistir
  3. Hicbiri yoksa yükleme talimatı yaz
"""
import sys
import os
import subprocess
from pathlib import Path


def md_to_docx_via_python_docx(src: Path, dst: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False

    doc = Document()
    
    # Sayfa kenar boslugu
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    current_code_block = []
    in_code_block = False

    def flush_code_block(lines):
        if not lines:
            return
        para = doc.add_paragraph()
        para.style = 'No Spacing'
        run = para.add_run('\n'.join(lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

    lines = src.read_text(encoding='utf-8').splitlines()

    for line in lines:
        # Kod blogu
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code_block(current_code_block)
                current_code_block = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            current_code_block.append(line)
            continue

        # Basliklar
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:], style='Quote')
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif line.startswith('| ') or line.startswith('|-'):
            # Tablo satirlarini basit metin olarak ekle
            p = doc.add_paragraph(line, style='No Spacing')
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('  - ') or line.startswith('  * '):
            doc.add_paragraph(line[4:], style='List Bullet 2')
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            # Normal paragraf — markdown kalin/italik temizle
            clean = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean)

    doc.save(str(dst))
    return True


def md_to_docx_via_pandoc(src: Path, dst: Path) -> bool:
    try:
        result = subprocess.run(
            ['pandoc', str(src), '-o', str(dst), '--from=markdown', '--to=docx'],
            capture_output=True, text=True, timeout=60
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def main():
    if len(sys.argv) < 2:
        print("Kullanim: python scripts/md_to_docx.py <kaynak.md> [hedef.docx]")
        sys.exit(1)

    src = Path(sys.argv[1])
    if not src.exists():
        print(f"HATA: Kaynak dosya bulunamadi: {src}")
        sys.exit(1)

    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix('.docx')

    print(f"Donusturuluyor: {src} -> {dst}")

    if md_to_docx_via_pandoc(src, dst):
        print(f"[OK] pandoc ile uretildi: {dst}")
        return

    if md_to_docx_via_python_docx(src, dst):
        print(f"[OK] python-docx ile uretildi: {dst}")
        return

    print("[HATA] Ne pandoc ne python-docx mevcut.")
    print("Cozum secenekleri:")
    print("  1. pip install python-docx")
    print("  2. https://pandoc.org/installing.html adresinden pandoc kur")
    sys.exit(1)


if __name__ == '__main__':
    main()
